#!/usr/bin/env python3
"""
Tests for the Test Case Generator module
"""

import pytest
import tempfile
import os
from test_case_generator import TestCaseGenerator


class TestTestCaseGenerator:
    
    @pytest.fixture
    def sample_parsed_features(self):
        """Sample parsed features for testing"""
        return {
            'REDE_CARD': {
                'provider': 'REDE',
                'payment_method': 'CARD',
                'features': {
                    'Country': 'Brazil',
                    'Verify': 'TRUE',
                    'Authorize': 'TRUE',
                    'Capture': 'TRUE',
                    'Refund': 'FALSE',
                    'Currency': 'BRL',
                    'Sandbox': 'TRUE'
                }
            },
            'PAGARME_CARD': {
                'provider': 'PAGARME',
                'payment_method': 'CARD',
                'features': {
                    'Country': 'Brazil',
                    'Verify': 'TRUE',
                    'Authorize': 'IMPLEMENTED',
                    'Capture': '',  # Empty value
                    'Refund': 'TRUE',
                    '3DS': 'TRUE',
                    'Webhook': 'FALSE'
                }
            }
        }
    
    @pytest.fixture
    def generator_en(self):
        """English test case generator"""
        return TestCaseGenerator(locale='en')
    
    @pytest.fixture
    def generator_es(self):
        """Spanish test case generator"""
        return TestCaseGenerator(locale='es')
    
    def test_generator_initialization(self):
        """Test generator initialization with different locales"""
        # Test English
        gen_en = TestCaseGenerator(locale='en')
        assert gen_en.locale == 'en'
        assert gen_en.i18n is not None
        
        # Test Spanish
        gen_es = TestCaseGenerator(locale='es')
        assert gen_es.locale == 'es'
        
        # Test default
        gen_default = TestCaseGenerator()
        assert gen_default.locale == 'en'
    
    def test_is_feature_implemented(self, generator_en):
        """Test feature implementation detection"""
        # Test implemented values
        assert generator_en._is_feature_implemented('TRUE') == True
        assert generator_en._is_feature_implemented('IMPLEMENTED') == True
        assert generator_en._is_feature_implemented('YES') == True
        assert generator_en._is_feature_implemented('Y') == True
        assert generator_en._is_feature_implemented('1') == True
        assert generator_en._is_feature_implemented('SUPPORTED') == True
        assert generator_en._is_feature_implemented('AVAILABLE') == True
        
        # Test case insensitive
        assert generator_en._is_feature_implemented('true') == True
        assert generator_en._is_feature_implemented('True') == True
        assert generator_en._is_feature_implemented(' TRUE ') == True
        
        # Test not implemented values
        assert generator_en._is_feature_implemented('FALSE') == False
        assert generator_en._is_feature_implemented('NO') == False
        assert generator_en._is_feature_implemented('0') == False
        assert generator_en._is_feature_implemented('') == False
        assert generator_en._is_feature_implemented(None) == False
        assert generator_en._is_feature_implemented('RANDOM_VALUE') == False
        
        # Test values that are not boolean flags (like 'Brazil', 'BRL')
        assert generator_en._is_feature_implemented('Brazil') == False
        assert generator_en._is_feature_implemented('BRL') == False
    
    def test_generate_test_cases_for_features(self, generator_en, sample_parsed_features):
        """Test test case generation for features"""
        test_cases_data = generator_en.generate_test_cases_for_features(sample_parsed_features)
        
        # Should be a flat list of test cases
        assert isinstance(test_cases_data, list)
        assert len(test_cases_data) > 0
        
        # Separate master test cases from feature-specific test cases
        master_test_cases = [tc for tc in test_cases_data if tc['provider'] == 'All Providers']
        feature_test_cases = [tc for tc in test_cases_data if tc['provider'] != 'All Providers']
        
        # Should always have master test cases
        assert len(master_test_cases) > 0, "Master test cases should always be included"
        
        # Should have feature-specific test cases for the implemented features
        assert len(feature_test_cases) > 0, "Should have feature-specific test cases"
        
        # Check that all test cases have the expected table format columns
        for test_case in test_cases_data:
            assert 'id' in test_case
            assert 'provider' in test_case
            assert 'payment_method' in test_case
            assert 'description' in test_case
            assert 'passed' in test_case
            assert 'date' in test_case
            assert 'executer' in test_case
            assert 'evidence' in test_case
            
            # Check ID format - should be feature prefix + number + salt (e.g., CTR0001.abc123)
            assert '.' in test_case['id'], f"Test case ID should contain a salt: {test_case['id']}"
            base_id, salt = test_case['id'].split('.', 1)
            assert len(salt) == 6, f"Salt should be 6 characters long: {salt}"
            assert salt.isalnum(), f"Salt should be alphanumeric: {salt}"
        
        # Check master test cases
        for master_case in master_test_cases:
            assert master_case['provider'] == 'All Providers'
            assert master_case['payment_method'] == 'All Payment Methods'
            assert master_case['id'].startswith('MST'), f"Master test case ID should start with MST: {master_case['id']}"
        
        # Check feature-specific test cases
        for feature_case in feature_test_cases:
            assert feature_case['provider'] in ['REDE', 'PAGARME']
            assert feature_case['payment_method'] == 'CARD'
            
            # Check description no longer includes provider + payment method
            assert 'REDE' not in feature_case['description']
            assert 'PAGARME' not in feature_case['description']
    
    def test_generate_test_cases_no_implemented_features(self, generator_en):
        """Test when no features are implemented"""
        parsed_features = {
            'TEST_PROVIDER': {
                'provider': 'TEST',
                'payment_method': 'CARD',
                'features': {
                    'Country': 'FALSE',
                    'Verify': '',
                    'Authorize': 'NO'
                }
            }
        }
        
        test_cases_data = generator_en.generate_test_cases_for_features(parsed_features)
        
        # Separate master test cases from feature-specific test cases
        master_test_cases = [tc for tc in test_cases_data if tc['provider'] == 'All Providers']
        feature_test_cases = [tc for tc in test_cases_data if tc['provider'] != 'All Providers']
        
        # Should always have master test cases
        assert len(master_test_cases) > 0, "Master test cases should always be included"
        
        # Should have no feature-specific test cases since no features are implemented
        assert len(feature_test_cases) == 0, "Should have no feature-specific test cases when no features are implemented"
    
    def test_generate_markdown_document(self, generator_en, sample_parsed_features):
        """Test markdown document generation"""
        markdown_doc = generator_en.generate_markdown_document(
            sample_parsed_features,
            merchant_name="Test Merchant",
            include_metadata=True
        )
        
        # Check document structure
        assert "# Test Cases for Test Merchant" in markdown_doc
        assert "## Test Case Documentation" in markdown_doc
        assert "## Summary" in markdown_doc
        
        # Should have table format
        assert "| `ID` | Provider | Payment Method | Description | Passed | Date | Executer | Evidence |" in markdown_doc
        assert "|----|----------|----------------|-------------|--------|------|----------|----------|" in markdown_doc
        
        # Should have test cases in table rows
        assert "0001." in markdown_doc  # Test case IDs with salt format
        assert "| REDE |" in markdown_doc
        assert "| PAGARME |" in markdown_doc
        assert "| CARD |" in markdown_doc
        
        # Should have instructions
        assert "Fill in the 'Passed' column" in markdown_doc
    
    def test_generate_markdown_document_without_metadata(self, generator_en, sample_parsed_features):
        """Test markdown document generation without metadata"""
        markdown_doc = generator_en.generate_markdown_document(
            sample_parsed_features,
            merchant_name="Test Merchant",
            include_metadata=False
        )
        
        # Should not have metadata sections
        assert "Generated on:" not in markdown_doc
        assert "## Summary" not in markdown_doc
        
        # But should have test cases
        assert "## Test Case Documentation" in markdown_doc
        assert "| `ID` | Provider | Payment Method | Description | Passed | Date | Executer | Evidence |" in markdown_doc
        assert "0001." in markdown_doc  # Test case IDs with salt format
    
    def test_generate_summary_statistics(self, generator_en, sample_parsed_features):
        """Test summary statistics generation"""
        stats = generator_en.generate_summary_statistics(sample_parsed_features)
        
        # Check required fields
        assert 'total_providers' in stats
        assert 'total_test_cases' in stats
        assert 'total_implemented_features' in stats
        assert 'features_by_provider' in stats
        assert 'language' in stats
        
        # Check values
        assert stats['total_providers'] == 2
        assert stats['total_test_cases'] > 0
        assert stats['language'] == 'en'
        
        # Check features by provider
        assert 'REDE' in stats['features_by_provider']
        assert 'PAGARME' in stats['features_by_provider']
        assert stats['features_by_provider']['REDE'] > 0
        assert stats['features_by_provider']['PAGARME'] > 0
    
    def test_different_locales(self, sample_parsed_features):
        """Test test case generation in different languages"""
        # Test English
        gen_en = TestCaseGenerator(locale='en')
        test_cases_en = gen_en.generate_test_cases_for_features(sample_parsed_features)
        
        # Test Spanish
        gen_es = TestCaseGenerator(locale='es')
        test_cases_es = gen_es.generate_test_cases_for_features(sample_parsed_features)
        
        # Should have same number of test cases but different descriptions
        assert len(test_cases_en) == len(test_cases_es)
        
        # Get first test case from each
        if test_cases_en and test_cases_es:
            en_first = test_cases_en[0]
            es_first = test_cases_es[0]
            
            # Same base ID format (before the salt), but salt will be different
            en_base_id = en_first['id'].split('.')[0]
            es_base_id = es_first['id'].split('.')[0]
            assert en_base_id == es_base_id, "Base IDs should be the same"
            
            # Both should have salt format
            assert '.' in en_first['id'], "English ID should have salt format"
            assert '.' in es_first['id'], "Spanish ID should have salt format"
            
            # Descriptions may differ based on locale if translations are available
            assert en_first['description'] is not None
            assert es_first['description'] is not None
    
    def test_empty_parsed_features(self, generator_en):
        """Test with empty parsed features"""
        empty_features = {}
        
        test_cases_data = generator_en.generate_test_cases_for_features(empty_features)
        
        # Separate master test cases from feature-specific test cases
        master_test_cases = [tc for tc in test_cases_data if tc['provider'] == 'All Providers']
        feature_test_cases = [tc for tc in test_cases_data if tc['provider'] != 'All Providers']
        
        # Should always have master test cases
        assert len(master_test_cases) > 0, "Master test cases should always be included"
        
        # Should have no feature-specific test cases since no features are provided
        assert len(feature_test_cases) == 0, "Should have no feature-specific test cases when no features are provided"
    
    def test_markdown_document_structure(self, generator_en, sample_parsed_features):
        """Test specific markdown document structure requirements"""
        markdown_doc = generator_en.generate_markdown_document(sample_parsed_features)
        
        lines = markdown_doc.split('\n')
        
        # Should have proper heading structure
        h1_lines = [line for line in lines if line.startswith('# ')]
        h2_lines = [line for line in lines if line.startswith('## ')]
        h3_lines = [line for line in lines if line.startswith('### ')]
        
        assert len(h1_lines) >= 1  # Main title
        assert len(h2_lines) >= 2  # At least intro + providers
        assert len(h3_lines) >= 1  # Feature sections
        
        # Test cases should be one per line as requested - look for actual test case lines
        test_case_lines = [line for line in lines if line.startswith('**') and 
                          ('0001' in line or '0002' in line or '0003' in line)]  # Look for test case IDs
        for line in test_case_lines:
            # Each line should be a complete test case
            assert line.count('**') >= 2  # At least one bold section
            assert '(' in line and ')' in line  # Type in parentheses
            assert ':' in line  # Description separator
    
    def test_feature_value_edge_cases(self, generator_en):
        """Test edge cases in feature values"""
        edge_case_features = {
            'TEST_PROVIDER': {
                'provider': 'TEST',
                'payment_method': 'CARD',
                'features': {
                    'Feature1': '  TRUE  ',  # With spaces
                    'Feature2': 'true',      # Lowercase
                    'Feature3': 'True',      # Mixed case
                    'Feature4': 'IMPLEMENTED',
                    'Feature5': 'false',     # Should not be included
                    'Feature6': '',          # Empty
                    'Feature7': 'RANDOM',    # Unknown value
                }
            }
        }
        
        test_cases_data = generator_en.generate_test_cases_for_features(edge_case_features)
        
        # Separate master test cases from feature-specific test cases
        master_test_cases = [tc for tc in test_cases_data if tc['provider'] == 'All Providers']
        feature_test_cases = [tc for tc in test_cases_data if tc['provider'] != 'All Providers']
        
        # Should always have master test cases
        assert len(master_test_cases) > 0, "Master test cases should always be included"
        
        # Since the test features (Feature1, Feature2, etc.) don't exist in feature_rules.json,
        # there should be no feature-specific test cases for them
        assert len(feature_test_cases) == 0, "No test cases for unknown features"
    
    def test_generate_html_document(self, generator_en, sample_parsed_features):
        """Test HTML document generation"""
        html_doc = generator_en.generate_html_document(
            sample_parsed_features,
            merchant_name="Test Merchant",
            include_metadata=True
        )
        
        # Check HTML structure
        assert html_doc.startswith('<!DOCTYPE html>')
        assert '<html lang="en">' in html_doc
        assert '<title>Test Cases for Test Merchant</title>' in html_doc
        assert '</html>' in html_doc
        
        # Check CSS styling is included
        assert '<style>' in html_doc
        assert 'font-family: Arial' in html_doc
        
        # Check table structure
        assert '<table>' in html_doc
        assert '<th class="id-column">ID</th>' in html_doc
        assert '<th>Provider</th>' in html_doc
        assert '<th>Payment Method</th>' in html_doc
        assert '<th>Description</th>' in html_doc
        assert '<th>Passed</th>' in html_doc
        assert '<th>Date</th>' in html_doc
        assert '<th>Executer</th>' in html_doc
        assert '<th>Evidence</th>' in html_doc
        
        # Check content structure
        assert '<h1>Test Cases for Test Merchant</h1>' in html_doc
        assert '<h2>Test Case Documentation</h2>' in html_doc
        
        # Check test case data in table rows
        assert '<td>' in html_doc and '0001.' in html_doc  # Test case IDs with salt format
        assert '<td>REDE</td>' in html_doc
        assert '<td>PAGARME</td>' in html_doc
        assert '<td>CARD</td>' in html_doc
        
        # Check metadata section
        assert 'class="metadata"' in html_doc
        assert 'Generated on:' in html_doc
        
        # Check summary section
        assert 'class="summary"' in html_doc
        assert 'class="notes"' in html_doc
    
    def test_generate_html_document_without_metadata(self, generator_en, sample_parsed_features):
        """Test HTML document generation without metadata"""
        html_doc = generator_en.generate_html_document(
            sample_parsed_features,
            merchant_name="Test Merchant",
            include_metadata=False
        )
        
        # Should not have metadata sections
        assert 'Generated on:' not in html_doc
        assert 'class="summary"' not in html_doc
        assert 'class="metadata"' not in html_doc
        
        # But should have test cases and basic structure
        assert '<h2>Test Case Documentation</h2>' in html_doc
        assert '<table>' in html_doc
        assert '0001.' in html_doc  # Test case IDs with salt format
    
    def test_generate_docx_document(self, generator_en, sample_parsed_features):
        """Test DOCX document generation"""
        from docx.document import Document as DocxDocument
        from io import BytesIO
        
        doc = generator_en.generate_docx_document(
            sample_parsed_features,
            merchant_name="Test Merchant",
            include_metadata=True
        )
        
        # Check that it returns a Document object
        assert isinstance(doc, DocxDocument)
        
        # Save to BytesIO to verify it's a valid document
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Check that we can load it back
        from docx import Document
        loaded_doc = Document(doc_io)
        
        # Check document structure
        paragraphs = [p.text for p in loaded_doc.paragraphs]
        all_text = ' '.join(paragraphs)
        
        # Check title is present
        assert 'Test Cases for Test Merchant' in all_text
        
        # Check metadata is included
        assert 'Generated on:' in all_text
        assert 'Language: EN' in all_text
        assert 'Total Test Cases:' in all_text
        
        # Check introduction section
        assert 'Test Case Documentation' in all_text
        assert 'This document contains test cases' in all_text
        
        # Check that tables exist (DOCX saves tables separately from paragraphs)
        assert len(loaded_doc.tables) > 0
        
        # Check table structure
        table = loaded_doc.tables[0]
        header_row = table.rows[0]
        header_cells = [cell.text for cell in header_row.cells]
        
        expected_headers = ['ID', 'Provider', 'Payment Method', 'Description', 'Passed', 'Date', 'Executer', 'Evidence']
        assert header_cells == expected_headers
        
        # Check that test case rows exist
        assert len(table.rows) > 1  # More than just the header
        
        # Check test case content in table
        test_case_found = False
        for row in table.rows[1:]:  # Skip header row
            cells = [cell.text for cell in row.cells]
            # Check if ID has the new format (contains a dot for salt) and other expected values
            if '.' in cells[0] and cells[1] in ['REDE', 'PAGARME'] and cells[2] == 'CARD':
                test_case_found = True
                break
        assert test_case_found, "Expected test case content not found in table"
    
    def test_generate_docx_document_without_metadata(self, generator_en, sample_parsed_features):
        """Test DOCX document generation without metadata"""
        from docx.document import Document as DocxDocument
        
        doc = generator_en.generate_docx_document(
            sample_parsed_features,
            merchant_name="Test Merchant",
            include_metadata=False
        )
        
        # Check that it returns a Document object
        assert isinstance(doc, DocxDocument)
        
        # Check document content
        paragraphs = [p.text for p in doc.paragraphs]
        all_text = ' '.join(paragraphs)
        
        # Should not have metadata
        assert 'Generated on:' not in all_text
        assert 'Total Test Cases:' not in all_text
        
        # But should have basic structure
        assert 'Test Cases for Test Merchant' in all_text
        assert 'Test Case Documentation' in all_text 