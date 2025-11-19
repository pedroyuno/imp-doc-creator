#!/usr/bin/env python3
"""
Test Case Generator for Implementation Scope

This module generates merchant-facing test case documents based on parsed 
implementation scope. It creates documents in multiple formats (HTML, Markdown)
for easy import into Google Docs and other documentation platforms.
"""

import os
import json
import secrets
import string
from datetime import datetime
from typing import Dict, List, Any, Optional
from i18n_helper import I18nHelper
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement, qn
from jinja2 import Environment, FileSystemLoader, select_autoescape


class TestCaseGenerator:
    def __init__(self, locale: str = 'en', rules_file_paths: Optional[List[str]] = None):
        """
        Initialize the test case generator.
        
        Args:
            locale: Language locale for test case descriptions (en, es, pt)
            rules_file_paths: Optional list of rules file paths to load
        """
        self.locale = locale
        self.i18n = I18nHelper(default_locale=locale, rules_file_paths=rules_file_paths)
        
        # Set up Jinja2 environment for document templates
        template_dir = os.path.join(os.path.dirname(__file__), 'templates', 'documents')
        self.jinja_env = Environment(
            loader=FileSystemLoader(template_dir),
            autoescape=select_autoescape(['html', 'xml']),
            trim_blocks=True,
            lstrip_blocks=True
        )
    
    def _generate_test_case_salt(self, length: int = 6) -> str:
        """
        Generate a random salt for test case IDs.
        
        Args:
            length: Length of the salt string (default: 6)
            
        Returns:
            Random alphanumeric string of specified length
        """
        alphabet = string.ascii_lowercase + string.digits
        return ''.join(secrets.choice(alphabet) for _ in range(length))
    
    def _filter_test_cases_by_environment(self, test_cases: List[Dict], environment: str) -> List[Dict]:
        """
        Filter test cases based on the specified environment.
        
        Args:
            test_cases: List of test cases to filter
            environment: Environment to filter by ('sandbox', 'production', or 'both')
            
        Returns:
            Filtered list of test cases
        """
        if environment == 'both':
            return test_cases
        
        filtered_cases = []
        for test_case in test_cases:
            test_env = test_case.get('environment', 'both')
            if test_env == 'both' or test_env == environment:
                filtered_cases.append(test_case)
        
        return filtered_cases
    
    def _get_implemented_features(self, parsed_features: Dict[str, Any]) -> List[str]:
        """
        Extract list of implemented features from parsed data.
        
        Args:
            parsed_features: Dictionary of parsed features from CSV parser
            
        Returns:
            List of feature names that are implemented
        """
        implemented_features = []
        
        for provider_key, provider_data in parsed_features.items():
            features = provider_data['features']
            
            for feature_name, feature_value in features.items():
                if self._is_feature_implemented(feature_value) and feature_name not in implemented_features:
                    implemented_features.append(feature_name)
        
        return implemented_features
    
    def generate_integration_steps(self, parsed_features: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Generate integration steps for implemented features based on feature_rules.json.
        Includes master integration steps that apply to all conversions.
        
        Args:
            parsed_features: Dictionary of parsed features from CSV parser
            
        Returns:
            List of integration steps with step number, comment, and documentation URL
        """
        # Get implemented features
        implemented_features = self._get_implemented_features(parsed_features)
        
        # Get integration steps from rules in the order they appear in feature_rules.json
        integration_steps = []
        step_number = 1
        
        # ALWAYS include master integration steps first (regardless of implemented features)
        master_steps = self.i18n.get_master_integration_steps()
        for step in master_steps:
            integration_steps.append({
                'step_number': step_number,
                'feature_name': 'Master Rules',
                'comment': step['comment'],
                'documentation_url': step['documentation_url']
            })
            step_number += 1
        
        # Use the feature order from feature_rules.json (via i18n helper)
        feature_order = list(self.i18n.feature_rules.keys())
        
        # Track providers from parsed_features to include provider-specific steps
        providers_in_scope = set()
        for provider_payment_key, provider_data in parsed_features.items():
            providers_in_scope.add(provider_data['provider'])
        
        for feature_name in feature_order:
            if feature_name in implemented_features:
                # Get universal integration steps
                integration_steps_data = self.i18n.get_integration_steps_for_feature(feature_name)
                integration_step_list = integration_steps_data.get('universal', [])
                
                # Add universal steps
                if integration_step_list:
                    for step in integration_step_list:
                        integration_steps.append({
                            'step_number': step_number,
                            'feature_name': feature_name,
                            'comment': step['comment'],
                            'documentation_url': step['documentation_url']
                        })
                        step_number += 1
                
                # Add provider-specific steps for each provider in scope
                provider_specific_steps = integration_steps_data.get('provider_specific', {})
                for provider in providers_in_scope:
                    if provider in provider_specific_steps:
                        for step in provider_specific_steps[provider]:
                            integration_steps.append({
                                'step_number': step_number,
                                'feature_name': f"{feature_name} ({provider})",
                                'comment': step['comment'],
                                'documentation_url': step['documentation_url']
                            })
                            step_number += 1
        
        return integration_steps
        
    def generate_test_cases_for_features(self, parsed_features: Dict[str, Any], environment: str = 'both') -> List[Dict]:
        """
        Generate test cases for all implemented features from parsed data.
        Includes master test cases that apply to all conversions.
        
        Args:
            parsed_features: Dictionary of parsed features from CSV parser
            environment: Environment filter ('sandbox', 'production', or 'both')
            
        Returns:
            List of test cases in table format with provider, payment method, and test details
        """
        all_test_cases = []
        
        # ALWAYS include master test cases first (regardless of implemented features)
        master_test_cases = self.i18n.get_master_test_cases(self.locale)
        filtered_master_cases = self._filter_test_cases_by_environment(master_test_cases, environment)
        
        for test_case in filtered_master_cases:
            # Use original ID from feature_rules.json with random salt
            original_id = test_case['id']
            salt = self._generate_test_case_salt()
            
            table_test_case = {
                'id': f"{original_id}.{salt}",
                'provider': 'All Providers',
                'payment_method': 'All Payment Methods',
                'description': test_case['description'],
                'passed': '',  # Empty field for manual completion
                'date': '',    # Empty field for manual completion
                'executer': '',  # Empty field for manual completion
                'evidence': '',   # Empty field for manual completion
                'environment': test_case.get('environment', 'both')
            }
            all_test_cases.append(table_test_case)
        
        # Generate test cases for implemented features
        for provider_payment_key, provider_data in parsed_features.items():
            provider = provider_data['provider']
            payment_method = provider_data['payment_method']
            features = provider_data['features']
            
            # Generate test cases for implemented features only
            for feature_name, feature_value in features.items():
                # Only include test cases for implemented features
                if self._is_feature_implemented(feature_value):
                    feature_test_cases = self.i18n.get_test_cases_for_feature(
                        feature_name,
                        self.locale,
                        payment_method=payment_method
                    )
                    
                    # Filter test cases by environment
                    filtered_test_cases = self._filter_test_cases_by_environment(feature_test_cases, environment)
                    
                    # Convert each test case to table format
                    for test_case in filtered_test_cases:
                        # Use original ID from feature_rules.json with random salt
                        original_id = test_case['id']
                        salt = self._generate_test_case_salt()
                        
                        table_test_case = {
                            'id': f"{original_id}.{salt}",
                            'provider': provider,
                            'payment_method': payment_method,
                            'description': test_case['description'],  # Remove provider + payment method from description
                            'passed': '',  # Empty field for manual completion
                            'date': '',    # Empty field for manual completion
                            'executer': '',  # Empty field for manual completion
                            'evidence': '',   # Empty field for manual completion
                            'environment': test_case.get('environment', 'both')  # Include environment info
                        }
                        all_test_cases.append(table_test_case)
        
        return all_test_cases
    
    def generate_environment_separated_test_cases(self, parsed_features: Dict[str, Any]) -> Dict[str, List[Dict]]:
        """
        Generate test cases separated by environment (sandbox and production).
        
        Args:
            parsed_features: Dictionary of parsed features from CSV parser
            
        Returns:
            Dictionary with 'sandbox' and 'production' keys containing their respective test cases
        """
        sandbox_cases = self.generate_test_cases_for_features(parsed_features, 'sandbox')
        production_cases = self.generate_test_cases_for_features(parsed_features, 'production')
        
        return {
            'sandbox': sandbox_cases,
            'production': production_cases
        }
    
    def generate_markdown_document(self, parsed_features: Dict[str, Any], 
                                 merchant_name: str = "Merchant", 
                                 include_metadata: bool = True,
                                 environment: str = 'both') -> str:
        """
        Generate a complete markdown document with test cases in table format.
        
        Args:
            parsed_features: Dictionary of parsed features from CSV parser
            merchant_name: Name of the merchant for document header
            include_metadata: Whether to include document metadata
            environment: Environment filter ('sandbox', 'production', 'both', or 'separated')
            
        Returns:
            Complete markdown document as string with table format
        """
        # Prepare template context
        context = {
            'merchant_name': merchant_name,
            'include_metadata': include_metadata,
            'environment': environment,
            'language': self.locale,
            'generation_timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'integration_steps': self.generate_integration_steps(parsed_features)
        }
        
        if environment == 'separated':
            # Generate separate tables for sandbox and production
            env_test_cases = self.generate_environment_separated_test_cases(parsed_features)
            context['env_test_cases_data'] = env_test_cases
            context['statistics'] = {
                'total_test_cases': len(env_test_cases['sandbox']) + len(env_test_cases['production']),
                'sandbox_test_cases': len(env_test_cases['sandbox']),
                'production_test_cases': len(env_test_cases['production'])
            }
            
            template = self.jinja_env.get_template('test_case_table_separated.md')
        else:
            # Generate single table for specified environment
            test_cases_data = self.generate_test_cases_for_features(parsed_features, environment)
            context['test_cases_data'] = test_cases_data
            context['statistics'] = {
                'total_test_cases': len(test_cases_data)
            }
            
            template = self.jinja_env.get_template('test_case_table_single.md')
        
        return template.render(context)
    
    def generate_html_document(self, parsed_features: Dict[str, Any], 
                             merchant_name: str = "Merchant", 
                             include_metadata: bool = True,
                             environment: str = 'both') -> str:
        """
        Generate a complete HTML document with test cases in table format.
        Google Docs can import HTML files directly while preserving formatting.
        
        Args:
            parsed_features: Dictionary of parsed features from CSV parser
            merchant_name: Name of the merchant for document header
            include_metadata: Whether to include document metadata
            environment: Environment filter ('sandbox', 'production', 'both', or 'separated')
            
        Returns:
            Complete HTML document as string with table format
        """
        # Prepare template context
        context = {
            'merchant_name': merchant_name,
            'include_metadata': include_metadata,
            'environment': environment,
            'language': self.locale,
            'generation_timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'integration_steps': self.generate_integration_steps(parsed_features)
        }
        
        if environment == 'separated':
            # Generate separate tables for sandbox and production
            env_test_cases = self.generate_environment_separated_test_cases(parsed_features)
            context['env_test_cases_data'] = env_test_cases
            context['statistics'] = {
                'total_test_cases': len(env_test_cases['sandbox']) + len(env_test_cases['production']),
                'sandbox_test_cases': len(env_test_cases['sandbox']),
                'production_test_cases': len(env_test_cases['production'])
            }
            
            template = self.jinja_env.get_template('test_case_table_separated.html')
        else:
            # Generate single table for specified environment
            test_cases_data = self.generate_test_cases_for_features(parsed_features, environment)
            context['test_cases_data'] = test_cases_data
            context['statistics'] = {
                'total_test_cases': len(test_cases_data)
            }
            
            template = self.jinja_env.get_template('test_case_table_single.html')
        
        return template.render(context)
    
    def generate_docx_document(self, parsed_features: Dict[str, Any], 
                             merchant_name: str = "Merchant", 
                             include_metadata: bool = True,
                             environment: str = 'both') -> Document:
        """
        Generate a complete DOCX document with test cases in table format.
        DOCX format provides professional styling and is compatible with Microsoft Word.
        
        Args:
            parsed_features: Dictionary of parsed features from CSV parser
            merchant_name: Name of the merchant for document header
            include_metadata: Whether to include document metadata
            environment: Environment filter ('sandbox', 'production', 'both', or 'separated')
            
        Returns:
            python-docx Document object ready for saving
        """
        # Create a new document
        doc = Document()
        
        # Set page orientation to landscape and adjust margins for better table fit
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)    # Standard letter height becomes width in landscape
        section.page_height = Inches(8.5)  # Standard letter width becomes height in landscape
        
        # Set margins for better table fit
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        # Document header
        if include_metadata:
            title = doc.add_heading(f'Test Cases for {merchant_name}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata paragraph
            metadata_para = doc.add_paragraph()
            metadata_para.add_run(f"Generated on: ").bold = True
            metadata_para.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            metadata_para.add_run(f"Language: ").bold = True
            metadata_para.add_run(f"{self.locale.upper()}\n")
            metadata_para.add_run(f"Environment: ").bold = True
            metadata_para.add_run(f"{environment.title()}")
        else:
            title = doc.add_heading(f'Test Cases for {merchant_name}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a line break
        doc.add_paragraph()
        
        # Integration Steps Section
        integration_steps = self.generate_integration_steps(parsed_features)
        if integration_steps:
            steps_heading = doc.add_heading('Integration Steps', level=1)
            steps_intro = doc.add_paragraph(
                'Follow these integration steps in sequential order to implement the required features for your payment integration:'
            )
            
            for step in integration_steps:
                step_heading = doc.add_heading(f"Step {step['step_number']}: {step['feature_name']} Implementation", level=2)
                
                desc_para = doc.add_paragraph()
                desc_para.add_run("Description: ").bold = True
                desc_para.add_run(step['comment'])
                
                doc_para = doc.add_paragraph()
                doc_para.add_run("Documentation: ").bold = True
                # Add clickable hyperlink
                self._add_hyperlink(doc_para, step['documentation_url'], step['documentation_url'])
        
        # Introduction section
        intro_heading = doc.add_heading('Test Case Documentation', level=1)
        intro_para = doc.add_paragraph(
            'This document contains test cases for your payment integration implementation. '
            'Each test case should be executed to ensure proper functionality of the implemented features.'
        )
        
        if environment == 'separated':
            # Generate separate tables for sandbox and production
            env_test_cases = self.generate_environment_separated_test_cases(parsed_features)
            
            for env_name in ['sandbox', 'production']:
                test_cases_data = env_test_cases[env_name]
                
                if test_cases_data:  # Only show section if there are test cases
                    # Add environment section header
                    env_heading = doc.add_heading(f'{env_name.title()} Environment Test Cases', level=2)
                    doc.add_paragraph(f'Test cases specifically for the {env_name} environment.')
                    
                    # Create table for this environment
                    self._create_docx_table(doc, test_cases_data, env_name)
        else:
            # Generate single table for specified environment
            test_cases_data = self.generate_test_cases_for_features(parsed_features, environment)
            self._create_docx_table(doc, test_cases_data)
        
        # Summary section
        if include_metadata:
            if environment == 'separated':
                env_test_cases = self.generate_environment_separated_test_cases(parsed_features)
                total_cases = len(env_test_cases['sandbox']) + len(env_test_cases['production'])
                sandbox_count = len(env_test_cases['sandbox'])
                production_count = len(env_test_cases['production'])
            else:
                test_cases_data = self.generate_test_cases_for_features(parsed_features, environment)
                total_cases = len(test_cases_data)
                sandbox_count = production_count = None
                
            doc.add_page_break()
            summary_heading = doc.add_heading('Summary', level=1)
            
            summary_para = doc.add_paragraph()
            summary_para.add_run("Total Test Cases: ").bold = True
            summary_para.add_run(f"{total_cases}\n")
            summary_para.add_run("Document Language: ").bold = True
            summary_para.add_run(f"{self.locale.upper()}\n")
            summary_para.add_run("Environment Filter: ").bold = True
            summary_para.add_run(f"{environment.title()}")
            
            if environment == 'separated':
                summary_para.add_run("\nSandbox Test Cases: ").bold = True
                summary_para.add_run(f"{sandbox_count}\n")
                summary_para.add_run("Production Test Cases: ").bold = True
                summary_para.add_run(f"{production_count}")
            
            # Instructions section
            instructions_heading = doc.add_heading('Instructions', level=2)
            instructions = [
                'Fill in the "Passed" column with Yes/No after executing each test case',
                'Record the execution date in the "Date" column',
                'Add the name of the person who executed the test in the "Executer" column',
                'Provide evidence (screenshots, logs, etc.) in the "Evidence" column'
            ]
            
            for instruction in instructions:
                para = doc.add_paragraph(instruction, style='List Bullet')
            
            # Footer note
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_run = footer_para.add_run(
                'This document was automatically generated from your implementation scope.'
            )
            footer_run.italic = True
        
        return doc
    
    def _create_docx_table(self, doc, test_cases_data, environment_name=None):
        """
        Create a DOCX table for test cases.
        
        Args:
            doc: Document object
            test_cases_data: List of test cases
            environment_name: Optional environment name for styling
        """
        # Create table with explicit width settings
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        
        # Set table to use fixed column widths
        table.autofit = False
        table.allow_autofit = False
        
        # Apply column width settings at XML level
        self._set_table_column_widths(table)
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Provider'
        hdr_cells[2].text = 'Payment Method'
        hdr_cells[3].text = 'Description'
        hdr_cells[4].text = 'Passed'
        hdr_cells[5].text = 'Date'
        hdr_cells[6].text = 'Executer'
        hdr_cells[7].text = 'Evidence'
        
        # Format header row with styling based on environment
        header_color = "366092"  # Default blue
        if environment_name == 'sandbox':
            header_color = "f39c12"  # Orange
        elif environment_name == 'production':
            header_color = "e74c3c"  # Red
        
        for i, cell in enumerate(hdr_cells):
            # Set background color
            self._set_cell_background_color(cell, header_color)
            
            # Make text bold, white, and centered
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                    run.font.size = Inches(0.11)  # Slightly larger font
                    if i == 0:  # ID column header - apply monospace font
                        run.font.name = 'Courier New'
        
        # Add test case rows with alternating colors
        for i, test_case in enumerate(test_cases_data):
            row_cells = table.add_row().cells
            row_cells[0].text = test_case['id']
            row_cells[1].text = test_case['provider']
            row_cells[2].text = test_case['payment_method']
            row_cells[3].text = test_case['description']
            row_cells[4].text = test_case['passed']
            row_cells[5].text = test_case['date']
            row_cells[6].text = test_case['executer']
            row_cells[7].text = test_case['evidence']
            
            # Add alternating row colors (light gray for even rows)
            if i % 2 == 0:
                for cell in row_cells:
                    self._set_cell_background_color(cell, "F8F9FA")  # Light gray
            
            # Apply monospace font to ID column and center align smaller columns
            for j, cell in enumerate(row_cells):
                if j == 0:  # ID column - apply monospace font
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.name = 'Courier New'
                elif j in [4, 5]:  # Passed, Date columns
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set column widths optimized for landscape orientation (total ~9.0")
        # Use table-level column width settings for better reliability
        table.columns[0].width = Inches(0.5)   # ID column (smallest)
        table.columns[1].width = Inches(0.8)   # Provider column (small)
        table.columns[2].width = Inches(1.0)   # Payment Method column (small)
        table.columns[3].width = Inches(4.8)   # Description column (LARGEST - main content)
        table.columns[4].width = Inches(0.6)   # Passed column (small)
        table.columns[5].width = Inches(0.7)   # Date column (small)
        table.columns[6].width = Inches(0.8)   # Executer column (small)
        table.columns[7].width = Inches(1.0)   # Evidence column (small)
        
        # Also set individual cell widths as backup
        for row in table.rows:
            row.cells[0].width = Inches(0.5)   # ID
            row.cells[1].width = Inches(0.8)   # Provider
            row.cells[2].width = Inches(1.0)   # Payment Method
            row.cells[3].width = Inches(4.8)   # Description (LARGEST)
            row.cells[4].width = Inches(0.6)   # Passed
            row.cells[5].width = Inches(0.7)   # Date
            row.cells[6].width = Inches(0.8)   # Executer
            row.cells[7].width = Inches(1.0)   # Evidence
    
    def generate_summary_statistics(self, parsed_features: Dict[str, Any], environment: str = 'both') -> Dict[str, Any]:
        """
        Generate summary statistics for the test cases.
        
        Args:
            parsed_features: Dictionary of parsed features from CSV parser
            environment: Environment filter ('sandbox', 'production', 'both', or 'separated')
            
        Returns:
            Dictionary containing summary statistics
        """
        if environment == 'separated':
            env_test_cases = self.generate_environment_separated_test_cases(parsed_features)
            sandbox_cases = env_test_cases['sandbox']
            production_cases = env_test_cases['production']
            total_test_cases = len(sandbox_cases) + len(production_cases)
        else:
            test_cases_data = self.generate_test_cases_for_features(parsed_features, environment)
            total_test_cases = len(test_cases_data)
            sandbox_cases = production_cases = None
        
        # Count implemented features per provider
        features_by_provider = {}
        total_implemented_features = 0
        
        for provider_key, provider_data in parsed_features.items():
            provider = provider_data['provider']
            implemented_count = 0
            
            for feature_name, feature_value in provider_data['features'].items():
                if self._is_feature_implemented(feature_value):
                    implemented_count += 1
            
            features_by_provider[provider] = implemented_count
            total_implemented_features += implemented_count
        
        stats = {
            'total_providers': len(parsed_features),
            'total_test_cases': total_test_cases,
            'total_implemented_features': total_implemented_features,
            'features_by_provider': features_by_provider,
            'language': self.locale,
            'environment': environment
        }
        
        if environment == 'separated':
            stats['sandbox_test_cases'] = len(sandbox_cases) if sandbox_cases else 0
            stats['production_test_cases'] = len(production_cases) if production_cases else 0
        
        return stats
    
    def _add_hyperlink(self, paragraph, url, text):
        """
        Add a clickable hyperlink to a paragraph.
        
        Args:
            paragraph: The paragraph to add the hyperlink to
            url: The URL to link to
            text: The text to display for the hyperlink
        """
        # Add relationship to document part
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        # Create hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # Create run element for the hyperlink text
        run = OxmlElement('w:r')
        r_pr = OxmlElement('w:rPr')
        
        # Set blue color
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')  # Blue
        r_pr.append(color)
        
        # Set underline
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')  # Underline
        r_pr.append(u)
        
        run.append(r_pr)
        
        # Add text
        text_elem = OxmlElement('w:t')
        text_elem.set(qn('xml:space'), 'preserve')  # Preserve spaces
        text_elem.text = text
        run.append(text_elem)
        
        hyperlink.append(run)
        
        # Append hyperlink to paragraph
        paragraph._p.append(hyperlink)
    
    def _set_cell_background_color(self, cell, color_hex):
        """Set the background color of a table cell.
        
        Args:
            cell: Table cell object
            color_hex: Hex color string (e.g., "366092" for blue)
        """
        # Create shading element using low-level XML manipulation
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex)
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    def _set_table_column_widths(self, table):
        """Force table to use specific column widths by setting XML properties."""
        # Define column widths in twips (1 inch = 1440 twips)
        column_widths = [
            720,   # 0.5 inches - ID
            1152,  # 0.8 inches - Provider
            1440,  # 1.0 inches - Payment Method
            6912,  # 4.8 inches - Description (LARGEST)
            864,   # 0.6 inches - Passed
            1008,  # 0.7 inches - Date
            1152,  # 0.8 inches - Executer
            1440   # 1.0 inches - Evidence
        ]
        
        # Set table properties for fixed layout
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        
        # Set table layout to fixed
        tbl_layout = OxmlElement('w:tblLayout')
        tbl_layout.set(qn('w:type'), 'fixed')
        tbl_pr.append(tbl_layout)
        
        # Set column widths
        tbl_grid = tbl.tblGrid
        for i, width in enumerate(column_widths):
            if i < len(tbl_grid.gridCol_lst):
                tbl_grid.gridCol_lst[i].w = width  # Pass as integer, not string
    
    def _is_feature_implemented(self, feature_value: str) -> bool:
        """
        Check if a feature is considered implemented based on its value.
        
        Args:
            feature_value: The value of the feature from CSV
            
        Returns:
            True if feature is implemented, False otherwise
        """
        if not feature_value:
            return False
        
        # Normalize value for comparison
        value = feature_value.strip().upper()
        
        # Consider these values as "implemented"
        implemented_values = ['TRUE', 'IMPLEMENTED', 'YES', 'Y', '1', 'SUPPORTED', 'AVAILABLE']
        
        return value in implemented_values


def main():
    """Example usage of the test case generator"""
    # Example parsed features data (this would normally come from CSV parser)
    example_features = {
        'REDE_CARD': {
            'provider': 'REDE',
            'payment_method': 'CARD',
            'features': {
                'Country': 'Brazil',
                'Verify': 'TRUE',
                'Authorize': 'TRUE',
                'Capture': 'TRUE',
                'Refund': 'TRUE',
                'Currency': 'BRL',
                'Sandbox': 'TRUE'
            }
        },
        'PAGARME_CARD': {
            'provider': 'PAGARME',
            'payment_method': 'CARD',
            'features': {
                'Country': 'Brazil',
                'Verify': 'TRUE',
                'Authorize': 'TRUE',
                'Capture': 'FALSE',
                'Refund': 'TRUE',
                '3DS': 'TRUE',
                'Webhook': 'TRUE'
            }
        }
    }
    
    # Generate test cases in English
    generator = TestCaseGenerator(locale='en')
    
    print("=== Test Case Generation Example ===")
    print()
    
    # Generate markdown document
    markdown_doc = generator.generate_markdown_document(
        example_features, 
        merchant_name="Example Merchant"
    )
    
    print("Generated Markdown Document:")
    print("=" * 50)
    print(markdown_doc)
    print("=" * 50)
    
    # Generate HTML document
    print("\n=== HTML Document Sample ===")
    html_doc = generator.generate_html_document(
        example_features,
        merchant_name="Example Merchant"
    )
    print("HTML document generated (showing first 500 characters):")
    print(html_doc[:500] + "...")
    
    # Generate summary statistics
    stats = generator.generate_summary_statistics(example_features)
    print("\nSummary Statistics:")
    print(f"Total Providers: {stats['total_providers']}")
    print(f"Total Test Cases: {stats['total_test_cases']}")
    print(f"Total Features: {stats['total_implemented_features']}")
    print(f"Features by Provider: {stats['features_by_provider']}")


if __name__ == "__main__":
    main() 